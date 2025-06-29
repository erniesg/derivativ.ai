"""
Document Export Service

Handles exporting generated documents to various formats (PDF, DOCX) with
student and teacher versions.
"""

import logging
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentExportService:
    """Service for exporting documents to various formats."""

    def __init__(self, output_directory: Optional[str] = None):
        """Initialize the export service.

        Args:
            output_directory: Directory to save exported files. Defaults to temp directory.
        """
        self.output_directory = Path(output_directory or tempfile.gettempdir())
        self.output_directory.mkdir(parents=True, exist_ok=True)

        # Initialize ReportLab styles
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Setup custom styles for PDF generation."""
        # Title style
        self.styles.add(
            ParagraphStyle(
                "CustomTitle",
                parent=self.styles["Title"],
                fontSize=18,
                spaceAfter=20,
                textColor=colors.darkblue,
                alignment=1,  # Center alignment
            )
        )

        # Section header style
        self.styles.add(
            ParagraphStyle(
                "SectionHeader",
                parent=self.styles["Heading2"],
                fontSize=14,
                spaceBefore=15,
                spaceAfter=10,
                textColor=colors.darkgreen,
                borderWidth=1,
                borderColor=colors.lightgrey,
                borderPadding=5,
                backColor=colors.lightgrey,
            )
        )

        # Question style
        self.styles.add(
            ParagraphStyle(
                "Question",
                parent=self.styles["Normal"],
                fontSize=11,
                spaceBefore=10,
                spaceAfter=5,
                leftIndent=20,
            )
        )

        # Answer style
        self.styles.add(
            ParagraphStyle(
                "Answer",
                parent=self.styles["Normal"],
                fontSize=10,
                spaceBefore=5,
                spaceAfter=10,
                leftIndent=40,
                textColor=colors.blue,
                fontName="Helvetica-Oblique",
            )
        )

    async def export_document(
        self,
        document: dict[str, Any],
        format_type: str,
        version: str = "student",
        store_in_r2: bool = False,
        r2_service=None,
    ) -> dict[str, Any]:
        """Export a document to a specific format with optional R2 storage.

        Args:
            document: Generated document data
            format_type: Export format ('pdf', 'docx', 'html', 'markdown')
            version: Document version ('student', 'teacher', 'combined')
            store_in_r2: Whether to store the exported file in R2
            r2_service: R2 storage service instance

        Returns:
            Dictionary with export result and file info
        """
        document_id = document.get("document_id", "unknown")
        title = document.get("title", "document").replace(" ", "_").lower()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        logger.info(f"📤 Exporting document {document_id} to {format_type} ({version})")

        try:
            if format_type.lower() == "pdf":
                content, file_path = await self._export_single_pdf(document, title, version)
            elif format_type.lower() == "docx":
                content, file_path = await self._export_single_docx(document, title, version)
            elif format_type.lower() == "html":
                content, file_path = await self._export_html(document, title, version)
            elif format_type.lower() == "markdown":
                content, file_path = await self._export_markdown(document, title, version)
            else:
                raise ValueError(f"Unsupported export format: {format_type}")

            result = {
                "success": True,
                "format": format_type,
                "version": version,
                "content": content if format_type in ["html", "markdown"] else None,
                "local_file_path": str(file_path) if file_path else None,
                "file_size": len(content) if isinstance(content, (str, bytes)) else 0,
            }

            # Store in R2 if requested
            if store_in_r2 and r2_service:
                logger.info(f"📤 Storing {format_type} file in R2...")

                # Generate R2 file key with proper extension
                file_extension = self._get_file_extension(format_type)
                r2_file_key = r2_service.generate_file_key(
                    document_id=document_id,
                    document_type="document",
                    file_format=file_extension,
                    version=version,
                )

                # Prepare file content for upload
                if isinstance(content, str):
                    file_content = content.encode("utf-8")
                elif isinstance(content, bytes):
                    file_content = content
                elif file_path:
                    # Read from file
                    with open(file_path, "rb") as f:
                        file_content = f.read()
                else:
                    raise ValueError("No content available for R2 upload")

                # Upload to R2
                upload_result = await r2_service.upload_file(
                    file_content=file_content,
                    file_key=r2_file_key,
                    metadata={
                        "document_id": document_id,
                        "format": format_type,
                        "version": version,
                        "title": document.get("title", ""),
                        "exported_at": datetime.now().isoformat(),
                        "content_type": self._get_content_type(format_type),
                    },
                )

                if upload_result["success"]:
                    result["r2_file_key"] = r2_file_key
                    result["r2_upload_id"] = upload_result["upload_id"]
                    
                    # Generate presigned URL for access (24 hour expiration)
                    try:
                        presigned_url = await r2_service.generate_presigned_url(
                            r2_file_key, expiration=86400  # 24 hours
                        )
                        result["r2_presigned_url"] = presigned_url
                        logger.info(f"✅ Stored in R2 with presigned URL: {r2_file_key}")
                    except Exception as e:
                        logger.warning(f"Failed to generate presigned URL: {e}")
                        result["r2_presigned_url"] = None
                        
                    logger.info(f"✅ Stored in R2: {r2_file_key}")
                else:
                    logger.error(f"❌ Failed to store in R2: {upload_result}")

            return result

        except Exception as e:
            logger.error(f"❌ Export failed: {e}")
            return {"success": False, "error": str(e), "format": format_type, "version": version}

    def _get_content_type(self, format_type: str) -> str:
        """Get MIME type for format."""
        content_types = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "html": "text/html",
            "markdown": "text/markdown",
            "txt": "text/plain",
        }
        return content_types.get(format_type.lower(), "application/octet-stream")

    def _get_file_extension(self, format_type: str) -> str:
        """Get proper file extension for format."""
        extension_map = {
            "pdf": "pdf",
            "docx": "docx",
            "html": "html",
            "markdown": "md",  # Use .md instead of .markdown
            "txt": "txt",
        }
        return extension_map.get(format_type.lower(), format_type.lower())

    async def _export_single_pdf(
        self, document: dict[str, Any], title: str, version: str
    ) -> tuple[bytes, Path]:
        """Export single PDF version."""
        output_path = self.output_directory / f"{title}_{version}.pdf"
        await self._create_pdf(document, output_path, version)

        # Read file content
        with open(output_path, "rb") as f:
            content = f.read()

        return content, output_path

    async def _export_single_docx(
        self, document: dict[str, Any], title: str, version: str
    ) -> tuple[bytes, Path]:
        """Export single DOCX version."""
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx not available for DOCX export")

        output_path = self.output_directory / f"{title}_{version}.docx"
        await self._create_docx(document, output_path, version)

        # Read file content
        with open(output_path, "rb") as f:
            content = f.read()

        return content, output_path

    async def _export_html(
        self, document: dict[str, Any], title: str, version: str
    ) -> tuple[str, None]:
        """Export to HTML format."""
        doc_title = document.get("title", "Generated Document")
        if version == "teacher":
            doc_title += " (Teacher Version)"
        elif version == "student":
            doc_title += " (Student Version)"

        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{doc_title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        h1 {{ color: #2c3e50; text-align: center; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #27ae60; margin-top: 30px; }}
        .metadata {{ background: #ecf0f1; padding: 15px; border-radius: 5px; margin: 20px 0; }}
        .question {{ background: #f8f9fa; padding: 15px; margin: 10px 0; border-left: 4px solid #3498db; }}
        .answer {{ background: #e8f5e8; padding: 10px; margin: 5px 0 15px 20px; border-left: 3px solid #27ae60; }}
        .objectives {{ background: #fff3cd; padding: 15px; border-radius: 5px; margin: 15px 0; }}
        ul {{ padding-left: 20px; }}
        .example {{ background: #f1f3f4; padding: 15px; margin: 15px 0; border-radius: 5px; }}
    </style>
</head>
<body>
    <h1>{doc_title}</h1>
"""

        # Add metadata
        metadata = self._extract_document_metadata(document)
        if metadata:
            html_content += '<div class="metadata"><h3>Document Information</h3>'
            for key, value in metadata.items():
                html_content += f"<p><strong>{key}:</strong> {value}</p>"
            html_content += "</div>"

        # Add content blocks
        content_structure = document.get("content_structure", {})
        blocks = content_structure.get("blocks", [])

        for block in blocks:
            html_content += await self._block_to_html(block, version)

        html_content += """
</body>
</html>"""

        return html_content, None

    async def _export_markdown(
        self, document: dict[str, Any], title: str, version: str
    ) -> tuple[str, None]:
        """Export to Markdown format."""
        doc_title = document.get("title", "Generated Document")
        if version == "teacher":
            doc_title += " (Teacher Version)"
        elif version == "student":
            doc_title += " (Student Version)"

        markdown_content = f"# {doc_title}\n\n"

        # Add metadata
        metadata = self._extract_document_metadata(document)
        if metadata:
            markdown_content += "## Document Information\n\n"
            for key, value in metadata.items():
                markdown_content += f"**{key}:** {value}  \n"
            markdown_content += "\n"

        # Add content blocks
        content_structure = document.get("content_structure", {})
        blocks = content_structure.get("blocks", [])

        for block in blocks:
            markdown_content += await self._block_to_markdown(block, version)

        return markdown_content, None

    async def _block_to_html(self, block: dict[str, Any], version: str) -> str:
        """Convert a content block to HTML."""
        block_type = block.get("block_type", "")
        content = block.get("content", {})
        block_title = self._format_block_title(block_type)

        html = f"<h2>{block_title}</h2>\n"

        if block_type == "learning_objectives":
            objectives = content.get("objectives", [])
            html += '<div class="objectives"><ul>\n'
            for objective in objectives:
                html += f"<li>{objective}</li>\n"
            html += "</ul></div>\n"

        elif block_type == "practice_questions":
            questions = content.get("questions", [])
            for i, question in enumerate(questions, 1):
                question_text = question.get("text", "")
                marks = question.get("marks", 0)

                html += '<div class="question">\n'
                html += f"<strong>Q{i}.</strong> {question_text}"
                if marks:
                    html += f" <em>[{marks} marks]</em>"
                html += "\n"

                if version in ["teacher", "combined"]:
                    answer = (
                        question.get("answer", "")
                        or question.get("solution", "")
                        or question.get("working", "")
                    )
                    if not answer:
                        answer = self._generate_basic_answer(question_text, i)
                    html += f'<div class="answer"><strong>Answer:</strong> {answer}</div>\n'

                html += "</div>\n"

        elif block_type == "worked_example":
            examples = content.get("examples", [])
            for i, example in enumerate(examples, 1):
                html += '<div class="example">\n'
                if isinstance(example, dict):
                    problem = example.get("problem", "")
                    answer = example.get("answer", "")
                    html += f"<strong>Example {i}:</strong> {problem}<br>\n"
                    if answer:
                        html += f"<strong>Answer:</strong> {answer}\n"
                else:
                    html += f"<strong>Example {i}:</strong> {example!s}\n"
                html += "</div>\n"

        else:
            # Generic content
            html += f"<div>{content!s}</div>\n"

        return html + "\n"

    async def _block_to_markdown(self, block: dict[str, Any], version: str) -> str:
        """Convert a content block to Markdown."""
        block_type = block.get("block_type", "")
        content = block.get("content", {})
        block_title = self._format_block_title(block_type)

        markdown = f"## {block_title}\n\n"

        if block_type == "learning_objectives":
            objectives = content.get("objectives", [])
            for objective in objectives:
                markdown += f"- {objective}\n"
            markdown += "\n"

        elif block_type == "practice_questions":
            questions = content.get("questions", [])
            for i, question in enumerate(questions, 1):
                question_text = question.get("text", "")
                marks = question.get("marks", 0)

                markdown += f"**Q{i}.** {question_text}"
                if marks:
                    markdown += f" *[{marks} marks]*"
                markdown += "\n\n"

                if version in ["teacher", "combined"]:
                    answer = (
                        question.get("answer", "")
                        or question.get("solution", "")
                        or question.get("working", "")
                    )
                    if not answer:
                        answer = self._generate_basic_answer(question_text, i)
                    markdown += f"> **Answer:** {answer}\n\n"

        elif block_type == "worked_example":
            examples = content.get("examples", [])
            for i, example in enumerate(examples, 1):
                if isinstance(example, dict):
                    problem = example.get("problem", "")
                    answer = example.get("answer", "")
                    markdown += f"**Example {i}:** {problem}\n\n"
                    if answer:
                        markdown += f"**Answer:** {answer}\n\n"
                else:
                    markdown += f"**Example {i}:** {example!s}\n\n"

        else:
            # Generic content
            markdown += f"{content!s}\n\n"

        return markdown

    async def _export_pdf(
        self, document: dict[str, Any], output_prefix: str, create_versions: bool
    ) -> list[str]:
        """Export document to PDF format."""
        exported_files = []

        if create_versions:
            # Create student version (no answers)
            student_file = self.output_directory / f"{output_prefix}_student.pdf"
            await self._create_pdf(document, student_file, version="student")
            exported_files.append(str(student_file))

            # Create teacher version (with answers)
            teacher_file = self.output_directory / f"{output_prefix}_teacher.pdf"
            await self._create_pdf(document, teacher_file, version="teacher")
            exported_files.append(str(teacher_file))
        else:
            # Create combined version
            combined_file = self.output_directory / f"{output_prefix}_combined.pdf"
            await self._create_pdf(document, combined_file, version="combined")
            exported_files.append(str(combined_file))

        return exported_files

    async def _create_pdf(self, document: dict[str, Any], output_path: Path, version: str):
        """Create a PDF file for the document."""
        # Create PDF document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=1 * inch,
            bottomMargin=1 * inch,
        )

        # Build content
        story = []

        # Add title
        title = document.get("title", "Generated Document")
        if version == "teacher":
            title += " (Teacher Version)"
        elif version == "student":
            title += " (Student Version)"

        story.append(Paragraph(title, self.styles["CustomTitle"]))
        story.append(Spacer(1, 20))

        # Add document metadata
        metadata = self._extract_document_metadata(document)
        if metadata:
            story.append(Paragraph("Document Information", self.styles["SectionHeader"]))
            for key, value in metadata.items():
                story.append(Paragraph(f"<b>{key}:</b> {value}", self.styles["Normal"]))
            story.append(Spacer(1, 15))

        # Add content blocks
        content_structure = document.get("content_structure", {})
        blocks = content_structure.get("blocks", [])

        for block in blocks:
            await self._add_block_to_pdf(story, block, version)

        # Build PDF
        doc.build(story)
        logger.info(f"Created PDF: {output_path}")

    async def _add_block_to_pdf(self, story: list, block: dict[str, Any], version: str):  # noqa: PLR0915
        """Add a content block to the PDF story."""
        block_type = block.get("block_type", "")
        content = block.get("content", {})

        # Add block header
        block_title = self._format_block_title(block_type)
        story.append(Paragraph(block_title, self.styles["SectionHeader"]))

        if block_type == "learning_objectives":
            objectives = content.get("objectives", [])
            for i, objective in enumerate(objectives, 1):
                story.append(Paragraph(f"{i}. {objective}", self.styles["Normal"]))
            story.append(Spacer(1, 10))

        elif block_type == "concept_explanation":
            concepts = content.get("concepts", [])
            for concept in concepts:
                name = concept.get("name", "")
                explanation = concept.get("explanation", "")
                story.append(Paragraph(f"<b>{name}</b>", self.styles["Normal"]))
                story.append(Paragraph(explanation, self.styles["Normal"]))
                story.append(Spacer(1, 10))

        elif block_type == "worked_example":
            examples = content.get("examples", [])
            for i, example in enumerate(examples, 1):
                # Handle both dict and string formats
                if isinstance(example, dict):
                    problem = example.get("problem", "")
                    steps = example.get("steps", [])
                    answer = example.get("answer", "")
                    explanation = example.get("explanation", "")

                    story.append(Paragraph(f"<b>Example {i}:</b> {problem}", self.styles["Normal"]))

                    if steps:
                        story.append(Paragraph("<b>Solution Steps:</b>", self.styles["Normal"]))
                        for step_num, step in enumerate(steps, 1):
                            step_text = step if isinstance(step, str) else str(step)
                            story.append(
                                Paragraph(f"{step_num}. {step_text}", self.styles["Normal"])
                            )

                    if answer:
                        story.append(Paragraph(f"<b>Answer:</b> {answer}", self.styles["Answer"]))

                    if explanation:
                        story.append(
                            Paragraph(f"<b>Explanation:</b> {explanation}", self.styles["Normal"])
                        )
                else:
                    # Handle string format
                    story.append(
                        Paragraph(f"<b>Example {i}:</b> {example!s}", self.styles["Normal"])
                    )

                story.append(Spacer(1, 15))

        elif block_type == "quick_reference":
            # Handle the raw dict format we're seeing
            if isinstance(content, dict):
                definitions = content.get("definitions", [])
                formulas = content.get("formulas", [])
                key_facts = content.get("key_facts", [])

                if definitions:
                    story.append(Paragraph("<b>Key Definitions:</b>", self.styles["Normal"]))
                    for defn in definitions:
                        if isinstance(defn, dict):
                            term = defn.get("term", "")
                            definition = defn.get("definition", "")
                            story.append(
                                Paragraph(f"• <b>{term}:</b> {definition}", self.styles["Normal"])
                            )
                        else:
                            story.append(Paragraph(f"• {defn!s}", self.styles["Normal"]))
                    story.append(Spacer(1, 10))

                if formulas:
                    story.append(Paragraph("<b>Key Formulas:</b>", self.styles["Normal"]))
                    for formula in formulas:
                        if isinstance(formula, dict):
                            name = formula.get("name", "")
                            expression = formula.get("expression", "")
                            story.append(
                                Paragraph(f"• <b>{name}:</b> {expression}", self.styles["Normal"])
                            )
                        else:
                            story.append(Paragraph(f"• {formula!s}", self.styles["Normal"]))
                    story.append(Spacer(1, 10))

                if key_facts:
                    story.append(Paragraph("<b>Key Facts:</b>", self.styles["Normal"]))
                    for fact in key_facts:
                        story.append(Paragraph(f"• {fact!s}", self.styles["Normal"]))
            else:
                # Fallback for other formats
                story.append(Paragraph(str(content), self.styles["Normal"]))
            story.append(Spacer(1, 10))

        elif block_type == "practice_questions":
            questions = content.get("questions", [])
            await self._add_questions_to_pdf(story, questions, version)

        elif block_type == "summary":
            key_points = content.get("key_points", [])
            for point in key_points:
                story.append(Paragraph(f"• {point}", self.styles["Normal"]))
            story.append(Spacer(1, 10))

        else:
            # Generic content handling
            story.append(Paragraph(str(content), self.styles["Normal"]))
            story.append(Spacer(1, 10))

    async def _add_questions_to_pdf(
        self, story: list, questions: list[dict[str, Any]], version: str
    ):
        """Add questions to PDF with appropriate formatting based on version."""
        for i, question in enumerate(questions, 1):
            # Question text
            question_text = question.get("text", "")
            marks = question.get("marks", 0)

            # Format question with marks
            formatted_question = f"Q{i}. {question_text}"
            if marks:
                formatted_question += f" [{marks} marks]"

            story.append(Paragraph(formatted_question, self.styles["Question"]))

            # Add answer space for student version
            if version == "student":
                story.append(Spacer(1, 30))  # Space for writing
                story.append(Paragraph("_" * 50, self.styles["Normal"]))
                story.append(Spacer(1, 20))

            # Add answers for teacher/combined version
            elif version in ["teacher", "combined"]:
                # Get answer - could be in different formats
                answer = question.get("answer", "")
                solution = question.get("solution", "")
                working = question.get("working", "")

                # Use the best available answer
                if solution:
                    story.append(Paragraph(f"<b>Solution:</b> {solution}", self.styles["Answer"]))
                elif answer:
                    story.append(Paragraph(f"<b>Answer:</b> {answer}", self.styles["Answer"]))
                elif working:
                    story.append(Paragraph(f"<b>Working:</b> {working}", self.styles["Answer"]))
                else:
                    # Generate a basic answer for the demo
                    basic_answer = self._generate_basic_answer(question_text, i)
                    story.append(Paragraph(f"<b>Answer:</b> {basic_answer}", self.styles["Answer"]))

                hint = question.get("hint", "")
                if hint:
                    story.append(Paragraph(f"<b>Hint:</b> {hint}", self.styles["Answer"]))

                story.append(Spacer(1, 15))

    def _generate_basic_answer(self, question_text: str, question_num: int) -> str:
        """Generate a basic answer for demonstration purposes."""
        # Simple pattern matching for common question types
        text_lower = question_text.lower()

        if "solve" in text_lower and "=" in question_text:
            return f"x = {question_num + 2}"
        elif "find" in text_lower and ("gradient" in text_lower or "slope" in text_lower):
            return f"Gradient = {question_num}"
        elif "calculate" in text_lower or "find" in text_lower:
            return f"Answer = {question_num * 5}"
        elif "equation" in text_lower and "line" in text_lower:
            return f"y = {question_num}x + {question_num * 2}"
        elif "graph" in text_lower or "plot" in text_lower:
            return "See graph with key points marked"
        else:
            return "Solution requires step-by-step working (see full solution guide)"

    async def _export_docx(
        self, document: dict[str, Any], output_prefix: str, create_versions: bool
    ) -> list[str]:
        """Export document to DOCX format."""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available, skipping DOCX export")
            return []

        exported_files = []

        if create_versions:
            # Create student version (no answers)
            student_file = self.output_directory / f"{output_prefix}_student.docx"
            await self._create_docx(document, student_file, version="student")
            exported_files.append(str(student_file))

            # Create teacher version (with answers)
            teacher_file = self.output_directory / f"{output_prefix}_teacher.docx"
            await self._create_docx(document, teacher_file, version="teacher")
            exported_files.append(str(teacher_file))
        else:
            # Create combined version
            combined_file = self.output_directory / f"{output_prefix}_combined.docx"
            await self._create_docx(document, combined_file, version="combined")
            exported_files.append(str(combined_file))

        return exported_files

    async def _create_docx(self, document: dict[str, Any], output_path: Path, version: str):
        """Create a DOCX file for the document."""
        doc = Document()

        # Add title
        title = document.get("title", "Generated Document")
        if version == "teacher":
            title += " (Teacher Version)"
        elif version == "student":
            title += " (Student Version)"

        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add document metadata
        metadata = self._extract_document_metadata(document)
        if metadata:
            doc.add_heading("Document Information", level=2)
            for key, value in metadata.items():
                p = doc.add_paragraph()
                run = p.add_run(f"{key}: ")
                run.bold = True
                p.add_run(value)

        # Add content blocks
        content_structure = document.get("content_structure", {})
        blocks = content_structure.get("blocks", [])

        for block in blocks:
            await self._add_block_to_docx(doc, block, version)

        # Save document
        doc.save(str(output_path))
        logger.info(f"Created DOCX: {output_path}")

    async def _add_block_to_docx(self, doc, block: dict[str, Any], version: str):  # noqa: PLR0915
        """Add a content block to the DOCX document."""
        block_type = block.get("block_type", "")
        content = block.get("content", {})

        # Add block header
        block_title = self._format_block_title(block_type)
        doc.add_heading(block_title, level=2)

        if block_type == "learning_objectives":
            objectives = content.get("objectives", [])
            for i, objective in enumerate(objectives, 1):
                doc.add_paragraph(f"{i}. {objective}")

        elif block_type == "concept_explanation":
            concepts = content.get("concepts", [])
            for concept in concepts:
                name = concept.get("name", "")
                explanation = concept.get("explanation", "")
                p = doc.add_paragraph()
                run = p.add_run(name)
                run.bold = True
                doc.add_paragraph(explanation)

        elif block_type == "worked_example":
            examples = content.get("examples", [])
            for i, example in enumerate(examples, 1):
                if isinstance(example, dict):
                    problem = example.get("problem", "")
                    steps = example.get("steps", [])
                    answer = example.get("answer", "")
                    explanation = example.get("explanation", "")

                    p = doc.add_paragraph()
                    run = p.add_run(f"Example {i}: ")
                    run.bold = True
                    p.add_run(problem)

                    if steps:
                        p = doc.add_paragraph()
                        run = p.add_run("Solution Steps:")
                        run.bold = True
                        for step_num, step in enumerate(steps, 1):
                            step_text = step if isinstance(step, str) else str(step)
                            doc.add_paragraph(f"{step_num}. {step_text}")

                    if answer:
                        p = doc.add_paragraph()
                        run = p.add_run("Answer: ")
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                        p.add_run(answer)

                    if explanation:
                        p = doc.add_paragraph()
                        run = p.add_run("Explanation: ")
                        run.bold = True
                        p.add_run(explanation)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(f"Example {i}: ")
                    run.bold = True
                    p.add_run(str(example))

        elif block_type == "quick_reference":
            if isinstance(content, dict):
                definitions = content.get("definitions", [])
                formulas = content.get("formulas", [])
                key_facts = content.get("key_facts", [])

                if definitions:
                    p = doc.add_paragraph()
                    run = p.add_run("Key Definitions:")
                    run.bold = True
                    for defn in definitions:
                        if isinstance(defn, dict):
                            term = defn.get("term", "")
                            definition = defn.get("definition", "")
                            p = doc.add_paragraph(f"• {term}: {definition}")
                        else:
                            doc.add_paragraph(f"• {defn!s}")

                if formulas:
                    p = doc.add_paragraph()
                    run = p.add_run("Key Formulas:")
                    run.bold = True
                    for formula in formulas:
                        if isinstance(formula, dict):
                            name = formula.get("name", "")
                            expression = formula.get("expression", "")
                            p = doc.add_paragraph(f"• {name}: {expression}")
                        else:
                            doc.add_paragraph(f"• {formula!s}")

                if key_facts:
                    p = doc.add_paragraph()
                    run = p.add_run("Key Facts:")
                    run.bold = True
                    for fact in key_facts:
                        doc.add_paragraph(f"• {fact!s}")
            else:
                doc.add_paragraph(str(content))

        elif block_type == "practice_questions":
            questions = content.get("questions", [])
            await self._add_questions_to_docx(doc, questions, version)

        elif block_type == "summary":
            key_points = content.get("key_points", [])
            for point in key_points:
                doc.add_paragraph(f"• {point}")
        else:
            # Generic content handling
            doc.add_paragraph(str(content))

    async def _add_questions_to_docx(self, doc, questions: list[dict[str, Any]], version: str):
        """Add questions to DOCX with appropriate formatting based on version."""
        for i, question in enumerate(questions, 1):
            # Question text
            question_text = question.get("text", "")
            marks = question.get("marks", 0)

            # Format question with marks
            formatted_question = f"Q{i}. {question_text}"
            if marks:
                formatted_question += f" [{marks} marks]"

            p = doc.add_paragraph(formatted_question)

            # Add answer space for student version
            if version == "student":
                doc.add_paragraph("")  # Empty line
                doc.add_paragraph("_" * 50)  # Answer line
                doc.add_paragraph("")  # Empty line

            # Add answers for teacher/combined version
            elif version in ["teacher", "combined"]:
                # Get answer - could be in different formats
                answer = question.get("answer", "")
                solution = question.get("solution", "")
                working = question.get("working", "")

                # Use the best available answer
                if solution:
                    p = doc.add_paragraph()
                    run = p.add_run("Solution: ")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                    p.add_run(solution)
                elif answer:
                    p = doc.add_paragraph()
                    run = p.add_run("Answer: ")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                    p.add_run(answer)
                elif working:
                    p = doc.add_paragraph()
                    run = p.add_run("Working: ")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                    p.add_run(working)
                else:
                    # Generate a basic answer for the demo
                    basic_answer = self._generate_basic_answer(question_text, i)
                    p = doc.add_paragraph()
                    run = p.add_run("Answer: ")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                    p.add_run(basic_answer)

                hint = question.get("hint", "")
                if hint:
                    p = doc.add_paragraph()
                    run = p.add_run("Hint: ")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                    p.add_run(hint)

                doc.add_paragraph("")  # Empty line

    def _extract_document_metadata(self, document: dict[str, Any]) -> dict[str, str]:
        """Extract metadata from document for display."""
        metadata = {}

        # Get generation request details
        generation_request = document.get("generation_request", {})
        if generation_request:
            topic = getattr(generation_request, "topic", None)
            if hasattr(topic, "value"):
                metadata["Topic"] = topic.value
            elif topic:
                metadata["Topic"] = str(topic)

            grade_level = getattr(generation_request, "grade_level", None)
            if grade_level:
                metadata["Grade Level"] = str(grade_level)

            difficulty = getattr(generation_request, "difficulty", None)
            if difficulty:
                metadata["Difficulty"] = str(difficulty)

        # Get document details
        estimated_minutes = document.get("total_estimated_minutes")
        if estimated_minutes:
            metadata["Estimated Duration"] = f"{estimated_minutes} minutes"

        detail_level = document.get("actual_detail_level")
        if detail_level:
            metadata["Detail Level"] = str(detail_level)

        return metadata

    def _format_block_title(self, block_type: str) -> str:
        """Format block type into a readable title."""
        title_map = {
            "learning_objectives": "Learning Objectives",
            "concept_explanation": "Concept Explanation",
            "worked_example": "Worked Examples",
            "practice_questions": "Practice Questions",
            "summary": "Summary",
            "extension_activities": "Extension Activities",
            "assessment_rubric": "Assessment Rubric",
        }

        return title_map.get(block_type, block_type.replace("_", " ").title())

    def get_export_directory(self) -> str:
        """Get the export directory path."""
        return str(self.output_directory)

    def cleanup_old_exports(self, max_age_hours: int = 24):
        """Clean up old exported files."""
        import time

        current_time = time.time()
        cutoff_time = current_time - (max_age_hours * 3600)

        for file_path in self.output_directory.glob("*.pdf"):
            if file_path.stat().st_mtime < cutoff_time:
                try:
                    file_path.unlink()
                    logger.info(f"Cleaned up old export: {file_path}")
                except Exception as e:
                    logger.warning(f"Failed to clean up {file_path}: {e}")
